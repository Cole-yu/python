#! /usr/bin/env python

#读取docx中的文本代码示例
import docx
from selenium import webdriver
import os
import time

#获取文档对象
file=docx.Document("D:\\desktop\\python.docx")
print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段

#输出每一段的内容
for para in file.paragraphs:
    print(para.text)

#输出段落编号及段落内容
for i in range(len(file.paragraphs)):
    print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)

txt=file.paragraphs[0].text
print(txt)

# url="http://win.gupiaotech.com/wap2/xbzg0424/index.html"
# driver = webdriver.Chrome()
# driver.get(url)
# driver.find_element_by_id("stock").send_keys("600120")
# driver.find_element_by_class_name("btn").click()
# time.sleep(5)
# driver.find_element_by_class_name("close").click()
# time.sleep(15)
# driver.quit()

url="https://weibo.com/"
# url="http://media.10jqka.com.cn/default/login/"

#浏览器启动，不显示提示“正受到自动测试工具控制”
option = webdriver.ChromeOptions()
# option.add_argument('disable-infobars')
# driver = webdriver.Chrome(chrome_options=option)

driver = webdriver.Chrome()
# driver = webdriver.Firefox() #模拟打开浏览器
driver.maximize_window()
driver.implicitly_wait(20)
driver.get(url)
driver.find_element_by_id("loginname").send_keys("13122102061")
driver.find_element_by_name("password").send_keys("qwer1234.")
driver.find_element_by_partial_link_text("登录").click()

first_window=driver.current_window_handle  #此行代码用来定位当前页面  
# driver.find_element_by_link_text("发布内容").click()

# # driver.quit()